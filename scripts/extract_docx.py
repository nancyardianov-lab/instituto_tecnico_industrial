#!/usr/bin/env python3
"""Extract text from all .docx files in /home/z/my-project/upload/"""
import os
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document

UPLOAD_DIR = Path("/home/z/my-project/upload")
OUT_DIR = Path("/home/z/my-project/scripts/extracted")
OUT_DIR.mkdir(parents=True, exist_ok=True)

for docx_path in sorted(UPLOAD_DIR.glob("*.docx")):
    print(f"\n{'='*80}")
    print(f"FILE: {docx_path.name}")
    print(f"{'='*80}")
    try:
        doc = Document(str(docx_path))
        # Paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[P{i}] {para.text}")
        # Tables
        for ti, table in enumerate(doc.tables):
            print(f"\n--- TABLE {ti} ---")
            for ri, row in enumerate(table.rows):
                cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
                print(f"  R{ri}: {cells}")
        # Images
        try:
            img_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    img_count += 1
            print(f"\n[IMAGES COUNT]: {img_count}")
        except Exception as e:
            print(f"[img count err]: {e}")

        # Also save extracted text to file
        out_file = OUT_DIR / (docx_path.stem + ".txt")
        with open(out_file, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + "\n")
            for ti, table in enumerate(doc.tables):
                f.write(f"\n--- TABLE {ti} ---\n")
                for row in table.rows:
                    cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
                    f.write("  " + " || ".join(cells) + "\n")
        print(f"\n[SAVED TO]: {out_file}")
    except Exception as e:
        print(f"ERROR reading {docx_path.name}: {e}")
