#!/usr/bin/env python3
"""Extract prompt.docx specifically including all tables."""
from docx import Document
from pathlib import Path

doc = Document("/home/z/my-project/upload/prompt.docx")

out = Path("/home/z/my-project/scripts/extracted/prompt.txt")
with open(out, "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPHS ===\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"[P{i}] {para.text}\n")
    
    f.write("\n=== TABLES ===\n")
    for ti, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {ti} ---\n")
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
            f.write(f"R{ri}: {cells}\n")

print(f"Saved to {out}")
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Show file
with open(out, "r", encoding="utf-8") as f:
    print(f.read())
